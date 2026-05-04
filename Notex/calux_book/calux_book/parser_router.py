"""Parser Router for Calux Book — lightweight, commercially-licensed document parsing.

All dependencies are permissively licensed (Apache-2.0, MIT, BSD):

**Fast-path (zero-AI, instant, all hardware):**
    .txt / .md / .log / code  → direct ``open()``
    .docx                     → python-docx  (~100 ms)
    .csv                      → direct read
    .xlsx / .xls / .ods       → python-calamine (Rust, ~50 ms)

**PDF pipeline — pypdfium2 + RapidOCR:**
    pypdfium2 text extraction (BSD-3 / Apache-2.0, ~50 ms per page)
    → quality check
    → RapidOCR image OCR fallback (Apache-2.0, ONNX-based)

**Images:**
    RapidOCR (Apache-2.0)
"""

from __future__ import annotations

import logging
import os
import re
from pathlib import Path
from typing import Any, Callable

logger = logging.getLogger("calux_book.parser_router")

# Type alias for progress callbacks
# callback(stage: str, detail: str, percent: int)
ProgressCallback = Callable[[str, str, int], None]

# ---------------------------------------------------------------------------
# Extension → parser mapping
# ---------------------------------------------------------------------------

_EXCEL_EXTS = {".xlsx", ".xls", ".xlsb", ".ods"}
_DOCX_EXTS = {".docx", ".doc"}
_CSV_EXTS = {".csv"}
_PDF_EXTS = {".pdf"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"}
_TEXT_EXTS = {".txt", ".log", ".json", ".xml", ".yaml", ".yml", ".toml",
              ".ini", ".cfg", ".rst", ".tex", ".py", ".js", ".ts", ".java",
              ".c", ".cpp", ".h", ".go", ".rs", ".sh", ".bat", ".ps1", ".md"}

# Minimum chars before we consider parsing "good enough"
_SPARSE_TEXT_THRESHOLD = 50

# Per-page minimum chars before triggering OCR fallback
_PAGE_MIN_CHARS = 15

# If more than 2% of alphabetic characters are replacement codepoints
# we treat the text layer as corrupted and fall back to OCR.
_CORRUPTION_RATIO_THRESHOLD = 0.02

# Text quality thresholds
_MIN_PRINTABLE_RATIO = 0.70
_MAX_CONTROL_RATIO = 0.03
_MAX_SYMBOL_RUNS = 8

# Common ligature codepoints that broken PDF fonts produce
_LIGATURE_MAP: dict[str, str] = {
    "\ufb00": "ff",
    "\ufb01": "fi",
    "\ufb02": "fl",
    "\ufb03": "ffi",
    "\ufb04": "ffl",
    "\ufb06": "st",
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

class ParserRouter:
    """Routes documents to the best parser based on file type.

    All parsers use commercially-licensed libraries only:
    - pypdfium2 (BSD-3/Apache-2.0) for PDF text extraction
    - RapidOCR (Apache-2.0) for image-based OCR fallback
    - python-docx (MIT) for DOCX
    - python-calamine (MIT) for Excel/ODS
    """

    def __init__(
        self,
        default_parser: str = "pdfium",
        complex_parser: str = "pdfium",
        ocr_fallback: str = "rapidocr",
        enable_ocr_fallback: bool = True,
        enable_fast_path: bool = True,
    ) -> None:
        self.default_parser = default_parser
        self.complex_parser = complex_parser
        self.ocr_fallback = ocr_fallback
        self.enable_ocr_fallback = enable_ocr_fallback
        self.enable_fast_path = enable_fast_path

        # Lazy-loaded parser instances
        self._rapid_ocr: Any = None

        # Progress callback (set externally per extraction)
        self._progress_cb: ProgressCallback | None = None

    # ------------------------------------------------------------------
    # Main entry point
    # ------------------------------------------------------------------

    def _notify(self, stage: str, detail: str, percent: int) -> None:
        """Safely invoke the progress callback if set."""
        if self._progress_cb:
            try:
                self._progress_cb(stage, detail, percent)
            except Exception:
                pass

    def extract(
        self, path: str, *, progress_cb: ProgressCallback | None = None,
    ) -> str:
        """Extract text from *path*, returning Markdown.

        Routing:
        1. Fast-path (zero-AI): text/docx/csv/excel
        2. PDF → pypdfium2 text extraction → RapidOCR fallback
        3. Images → RapidOCR
        """
        self._progress_cb = progress_cb
        ext = os.path.splitext(path)[1].lower()

        # ── Fast path (zero-AI) ──────────────────────────────────────
        if self.enable_fast_path:
            if ext in _TEXT_EXTS:
                self._notify("extracting", "Reading text file", 50)
                return self._read_text(path)
            if ext in _DOCX_EXTS:
                self._notify("extracting", "Parsing DOCX", 30)
                return self._parse_docx_fast(path)
            if ext in _CSV_EXTS:
                self._notify("extracting", "Reading CSV", 50)
                return self._read_text(path)
            if ext in _EXCEL_EXTS:
                self._notify("extracting", "Parsing spreadsheet", 30)
                return self._parse_excel(path)
        else:
            if ext in _TEXT_EXTS:
                return self._read_text(path)
            if ext in _EXCEL_EXTS:
                return self._parse_excel(path)

        # ── Images → OCR path ───────────────────────────────────────
        if ext in _IMAGE_EXTS:
            self._notify("extracting", "Running OCR on image", 20)
            return self._parse_image(path)

        # ── PDF → pypdfium2 + RapidOCR fallback ────────────────────
        if ext in _PDF_EXTS:
            return self._parse_pdf(path)

        # ── Unknown → try plain text, then binary check ─────────────
        return self._read_text_or_fail(path)

    def extract_from_url(self, url: str) -> str:
        """Fetch URL content via httpx and extract text."""
        try:
            import httpx
            resp = httpx.get(url, timeout=30, follow_redirects=True)
            resp.raise_for_status()
            content_type = resp.headers.get("content-type", "")

            if "pdf" in content_type:
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                    f.write(resp.content)
                    f.flush()
                    try:
                        return self._parse_pdf(f.name)
                    finally:
                        os.unlink(f.name)

            text = resp.text
            if text and text.strip():
                # Simple HTML tag stripping for plain text extraction
                if "<html" in text.lower() or "<body" in text.lower():
                    text = re.sub(r"<script[^>]*>.*?</script>", "", text,
                                  flags=re.DOTALL | re.IGNORECASE)
                    text = re.sub(r"<style[^>]*>.*?</style>", "", text,
                                  flags=re.DOTALL | re.IGNORECASE)
                    text = re.sub(r"<[^>]+>", " ", text)
                    text = re.sub(r"\s+", " ", text).strip()
                return text
            raise RuntimeError("URL returned empty content")
        except Exception as e:
            logger.error("URL extraction failed: %s", e)
            raise RuntimeError(f"Failed to fetch URL content: {e}") from e

    def extract_pages(
        self, path: str, *, progress_cb: ProgressCallback | None = None,
    ) -> list[dict[str, Any]]:
        """Extract page-aware content from *path*.

        Returns a list of dictionaries with keys:
            - page_number: int
            - text: str
            - section_title: str
            - block_type: str
        """
        self._progress_cb = progress_cb
        ext = os.path.splitext(path)[1].lower()

        if ext in _PDF_EXTS:
            self._notify("extracting", "Extracting PDF pages", 20)
            return self._extract_pdf_pages(path)

        text = self.extract(path, progress_cb=progress_cb)
        section = self._infer_section_title(text)
        return [{
            "page_number": 1,
            "text": text,
            "section_title": section,
            "block_type": self._infer_block_type(text),
        }]

    # ------------------------------------------------------------------
    # PDF: per-page hybrid (pypdfium2 text + RapidOCR OCR per page)
    # ------------------------------------------------------------------

    def _parse_pdf(self, path: str) -> str:
        """Extract text from PDF using per-page hybrid strategy.

        For each page: try pypdfium2 text → quality check → RapidOCR if
        needed.  No page-count cap.  Memory-efficient (one page at a time).
        """
        basename = os.path.basename(path)
        self._notify("extracting", f"Processing: {basename}", 10)

        pages = self._extract_pdf_pages_hybrid(path)

        if not pages:
            raise RuntimeError(
                f"Could not extract text from {basename}. "
                "The PDF may be image-only or contain no recognisable text."
            )

        out: list[str] = []
        for pg in pages:
            text = pg.get("text", "").strip()
            if text:
                out.append(f"[PAGE {pg['page_number']}]\n{text}")

        merged = "\n\n".join(out).strip()
        if not merged:
            raise RuntimeError(
                f"Could not extract text from {basename}. "
                "All pages returned empty text."
            )

        self._notify("extracting", "Text extraction complete", 85)
        logger.info(
            "PDF extraction done: %d pages with text from %s",
            len(pages), basename,
        )
        return merged

    def _extract_pdf_pages(self, path: str) -> list[dict[str, Any]]:
        """Extract page-aware content from PDF (delegates to hybrid)."""
        return self._extract_pdf_pages_hybrid(path)

    def _extract_pdf_pages_hybrid(self, path: str) -> list[dict[str, Any]]:
        """Per-page hybrid: text first, OCR fallback per page.

        For each page:
          1. Try pypdfium2 ``textpage.get_text_range()``
          2. If text < ``_PAGE_MIN_CHARS`` and OCR enabled → render + OCR
          3. Keep the page if it has any text; skip truly empty pages

        No page-count cap — processes the entire document.
        Memory-efficient: renders one page image at a time.
        """
        import pypdfium2 as pdfium

        basename = os.path.basename(path)
        pdf = pdfium.PdfDocument(path)
        total = len(pdf)
        pages: list[dict[str, Any]] = []
        ocr_engine: Any = None
        stats = {"text_ok": 0, "ocr_ok": 0, "skipped": 0, "corrupt": 0}

        # --- OCR-first auto-switch probe ---
        # If the first few pages all have corrupted text layers we skip
        # the (expensive) text→corruption→OCR cycle for every single page
        # and go straight to OCR for the remainder of the document.
        _PROBE_PAGES = 5
        _PROBE_THRESHOLD = 3
        ocr_first_mode = False

        logger.info(
            "Starting PDF extraction: %s (%d total pages)", basename, total,
        )

        try:
            for page_idx in range(total):
                page_no = page_idx + 1

                # Progress every 10 pages or on first/last
                if page_idx % 10 == 0 or page_idx == total - 1:
                    pct = int(20 + 60 * page_idx / max(total, 1))
                    self._notify(
                        "extracting", f"Page {page_no}/{total}", pct,
                    )

                text = ""
                used_ocr = False

                # --- Step 1: pypdfium2 text layer (skip if ocr_first_mode) ---
                if not ocr_first_mode:
                    try:
                        page_obj = pdf[page_idx]
                        textpage = page_obj.get_textpage()
                        raw = textpage.get_text_range()
                        if raw:
                            candidate = self._repair_ligatures(raw.strip())
                            if candidate:
                                # For substantial text, verify it isn't corrupted
                                if (
                                    len(candidate) >= _PAGE_MIN_CHARS
                                    and self._is_text_corrupted(candidate)
                                ):
                                    logger.info(
                                        "Page %d/%d text layer corrupted "
                                        "(%d chars), will try OCR — %s",
                                        page_no, total, len(candidate), basename,
                                    )
                                    stats["corrupt"] += 1
                                else:
                                    text = candidate
                    except Exception as exc:
                        logger.warning(
                            "Text layer failed page %d/%d of %s: %s",
                            page_no, total, basename, exc,
                        )

                # --- Probe-window auto-switch ---
                if (
                    not ocr_first_mode
                    and page_idx + 1 == min(_PROBE_PAGES, total)
                    and stats["corrupt"] >= _PROBE_THRESHOLD
                ):
                    ocr_first_mode = True
                    logger.warning(
                        "OCR-first mode activated after %d/%d probe "
                        "pages corrupted — %s (%d pages remain)",
                        stats["corrupt"], _PROBE_PAGES,
                        basename, total - page_idx - 1,
                    )

                # --- Step 2: OCR fallback if text too short ---
                if len(text) < _PAGE_MIN_CHARS and self.enable_ocr_fallback:
                    try:
                        if ocr_engine is None:
                            ocr_engine = self._get_rapidocr()

                        img = self._render_page_image(pdf[page_idx])
                        if img is not None:
                            import numpy as np

                            arr = np.array(img.convert("RGB"))
                            result = ocr_engine(arr)
                            txts = getattr(result, "txts", None)
                            if txts:
                                ocr_text = "\n".join(
                                    str(t).strip()
                                    for t in txts
                                    if str(t).strip()
                                )
                                if len(ocr_text) > len(text):
                                    text = ocr_text
                                    used_ocr = True
                    except Exception as exc:
                        logger.warning(
                            "OCR failed page %d/%d of %s: %s",
                            page_no, total, basename, exc,
                        )

                # --- Step 3: keep if non-empty ---
                text = text.strip()
                if text:
                    pages.append({
                        "page_number": page_no,
                        "text": text,
                        "section_title": self._infer_section_title(text),
                        "block_type": self._infer_block_type(text),
                    })
                    if used_ocr:
                        stats["ocr_ok"] += 1
                    else:
                        stats["text_ok"] += 1
                else:
                    stats["skipped"] += 1
                    if page_idx < 20 or page_idx % 50 == 0:
                        logger.warning(
                            "Page %d/%d empty after extraction+OCR — %s",
                            page_no, total, basename,
                        )
        finally:
            pdf.close()

        logger.info(
            "PDF extraction complete: %s — %d/%d pages extracted "
            "(text=%d, ocr=%d, corrupt_detected=%d, skipped=%d, "
            "ocr_first=%s)",
            basename, len(pages), total,
            stats["text_ok"], stats["ocr_ok"],
            stats["corrupt"], stats["skipped"],
            ocr_first_mode,
        )
        return pages

    # ------------------------------------------------------------------
    # Fast-path: python-docx (zero AI, ~100 ms)
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_docx_fast(path: str) -> str:
        """Extract text from a .docx using python-docx (MIT license).

        Inserts ``[PAGE N]`` markers at page breaks (``<w:br w:type="page"/>``)
        or every ~3000 chars at paragraph boundaries so non-PDF documents
        also benefit from page-wise chunking.
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise RuntimeError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            )

        doc = DocxDocument(path)
        paragraphs: list[str] = []
        page_markers: list[int] = []  # indices into paragraphs where a page break occurs

        for para in doc.paragraphs:
            text = para.text.strip()
            # Detect page breaks in the paragraph's XML runs
            has_page_break = False
            try:
                from lxml import etree
                for run in para._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
                ):
                    if run.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                    ) == "page":
                        has_page_break = True
                        break
            except Exception:
                pass

            if has_page_break and paragraphs:
                page_markers.append(len(paragraphs))
            if text:
                paragraphs.append(text)

        # Also grab tables
        for table in doc.tables:
            rows_md: list[str] = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows_md.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows_md.append(
                        "| " + " | ".join("---" for _ in cells) + " |"
                    )
            if rows_md:
                paragraphs.append("\n".join(rows_md))

        if not paragraphs:
            raise RuntimeError("DOCX parsed but no extractable text was found")

        # Build output with page markers
        # If we found explicit page breaks, use them
        if page_markers:
            out: list[str] = []
            page_num = 1
            prev_idx = 0
            for marker_idx in page_markers:
                page_text = "\n\n".join(paragraphs[prev_idx:marker_idx]).strip()
                if page_text:
                    out.append(f"[PAGE {page_num}]\n{page_text}")
                    page_num += 1
                prev_idx = marker_idx
            # Remaining paragraphs
            remaining = "\n\n".join(paragraphs[prev_idx:]).strip()
            if remaining:
                out.append(f"[PAGE {page_num}]\n{remaining}")
            return "\n\n".join(out).strip()

        # No explicit page breaks — insert synthetic markers every ~3000 chars
        content = "\n\n".join(paragraphs)
        return _insert_synthetic_page_markers(content)

    # ==================================================================
    #  Text quality helpers
    # ==================================================================

    @staticmethod
    def _repair_ligatures(text: str) -> str:
        """Replace known Unicode ligature codepoints with plain ASCII."""
        if not text:
            return text
        for lig, repl in _LIGATURE_MAP.items():
            text = text.replace(lig, repl)
        return text

    @staticmethod
    def _is_text_corrupted(text: str) -> bool:
        """Return *True* when the extracted text has too many broken characters.

        Designed to be lenient with non-Latin scripts (Hindi, Arabic, CJK, etc.)
        while still catching genuinely corrupted text layers.
        """
        if not text:
            return True
        stripped = text.strip()
        if len(stripped) < _SPARSE_TEXT_THRESHOLD:
            return True

        total = len(text)
        printable = sum(1 for ch in text if ch.isprintable() or ch in "\n\r\t")
        printable_ratio = printable / max(total, 1)

        alpha = 0
        bad = 0
        controls = 0
        for ch in text:
            if ch.isalpha():
                alpha += 1
            # Only count the Unicode REPLACEMENT CHARACTER as genuinely bad.
            # Extended-Latin, Devanagari, Arabic, CJK etc. are legitimate.
            if ch == "\ufffd":
                bad += 1
            if (not ch.isprintable()) and ch not in "\n\r\t":
                controls += 1

        control_ratio = controls / max(total, 1)
        if alpha == 0:
            # No alphabetic chars at all — but check for non-Latin letters
            # using a broader Unicode letter test
            unicode_letters = sum(1 for ch in text if ch.isalpha())
            if unicode_letters == 0:
                return True
        corruption_ratio = bad / max(alpha, 1)

        symbol_runs = len(re.findall(r"[^\w\s.,:;!?()\-]{8,}", text))
        # Use Unicode-aware letter matching: \w includes non-Latin letters
        # This prevents Hindi, Arabic, CJK pages from being flagged
        token_like = re.findall(r"\w{2,}", text)

        if printable_ratio < _MIN_PRINTABLE_RATIO:
            return True
        if control_ratio > _MAX_CONTROL_RATIO:
            return True
        if corruption_ratio > _CORRUPTION_RATIO_THRESHOLD:
            return True
        # Scale symbol-run tolerance by text length: longer pages (TOC,
        # indexes) naturally have more dot-leaders and separator runs.
        run_limit = max(_MAX_SYMBOL_RUNS, total // 200)
        if symbol_runs > run_limit:
            return True
        if len(token_like) < 3 and len(text) > 200:
            return True
        return False

    @staticmethod
    def _infer_section_title(text: str) -> str:
        lines = [ln.strip() for ln in text.splitlines()[:20] if ln.strip()]
        for ln in lines:
            if ln.startswith("#"):
                return ln.lstrip("# ").strip()[:160]
        if lines:
            return lines[0][:160]
        return ""

    @staticmethod
    def _infer_block_type(text: str) -> str:
        if "|" in text and "---" in text:
            return "table"
        return "text"

    def _accept_text(
        self, text: str | None, basename: str, parser: str,
    ) -> str | None:
        """Return cleaned text if long enough and not corrupted, else None."""
        if not text or len(text) <= _SPARSE_TEXT_THRESHOLD:
            logger.info(
                "%s: sparse (%d chars) for %s",
                parser, len(text) if text else 0, basename,
            )
            return None

        text = self._repair_ligatures(text)

        if self._is_text_corrupted(text):
            logger.warning(
                "%s: text corrupted for %s — will try OCR",
                parser, basename,
            )
            return None

        logger.debug(
            "%s: accepted %d chars from %s", parser, len(text), basename,
        )
        return text

    # ------------------------------------------------------------------
    # RapidOCR (fallback for scans / broken text layers)
    # ------------------------------------------------------------------

    def _get_rapidocr(self) -> Any:
        """Lazy-init RapidOCR with tuned thresholds.

        Compatible with RapidOCR v3.x which uses OmegaConf + enum
        types for lang_type.  We skip lang_type params to use the
        default Chinese models (which also handle English/Latin text
        well) and only override numeric thresholds.
        """
        if self._rapid_ocr is None:
            try:
                from rapidocr.main import RapidOCR
                # Only set scalar params — lang_type requires enum values
                # in v3.6+.  The default 'ch' model handles English fine.
                self._rapid_ocr = RapidOCR(params={
                    "Global.text_score": 0.3,
                    "Global.max_side_len": 4000,
                    "Det.box_thresh": 0.3,
                    "Det.unclip_ratio": 1.8,
                })
                logger.info("RapidOCR engine initialised (tuned thresholds)")
            except ImportError:
                raise RuntimeError(
                    "RapidOCR is required for OCR fallback. "
                    "Install it with: pip install rapidocr"
                )
        return self._rapid_ocr

    def _parse_rapidocr(self, path: str) -> str:
        """OCR a single image file using RapidOCR."""
        import numpy as np
        from PIL import Image

        img = Image.open(path).convert("RGB")
        arr = np.array(img)
        ocr = self._get_rapidocr()
        result = ocr(arr)
        txts = getattr(result, "txts", None)
        if not txts:
            return ""
        return "\n".join(str(t).strip() for t in txts if str(t).strip())

    def _parse_image(self, path: str) -> str:
        """Parse an image file using RapidOCR."""
        if self.enable_ocr_fallback:
            try:
                text = self._parse_rapidocr(path)
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(
                    "RapidOCR failed for image %s: %s",
                    os.path.basename(path), e,
                )

        raise RuntimeError(
            f"Could not extract text from image: {os.path.basename(path)}. "
            "Ensure RapidOCR is installed."
        )

    # ------------------------------------------------------------------
    # Excel via python-calamine
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_excel(path: str) -> str:
        """Parse Excel/ODS files using python-calamine (MIT, Rust reader)."""
        from python_calamine import CalamineWorkbook

        wb = CalamineWorkbook.from_path(path)
        lines: list[str] = []

        for sheet_name in wb.sheet_names:
            lines.append(f"## Sheet: {sheet_name}\n")
            rows = wb.get_sheet_by_name(sheet_name).to_python()
            if not rows:
                continue

            # First row as header
            header = rows[0]
            header_strs = [
                str(c).strip() if c is not None else "" for c in header
            ]
            lines.append("| " + " | ".join(header_strs) + " |")
            lines.append(
                "| " + " | ".join("---" for _ in header_strs) + " |"
            )

            for row in rows[1:]:
                cells = [
                    str(c).strip() if c is not None else "" for c in row
                ]
                if any(cells):
                    lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

        content = "\n".join(lines).strip()
        if not content:
            raise RuntimeError(
                "Excel file parsed but no extractable text was found"
            )
        return content

    # ------------------------------------------------------------------
    # Plain text
    # ------------------------------------------------------------------

    @staticmethod
    def _read_text(path: str) -> str:
        """Read plain text and insert synthetic [PAGE N] markers for large files."""
        data = Path(path).read_bytes()
        content = data.decode("utf-8", errors="replace")
        # Only insert page markers for files > 3000 chars
        if len(content) > 3000:
            return _insert_synthetic_page_markers(content)
        return content

    @staticmethod
    def _read_text_or_fail(path: str) -> str:
        data = Path(path).read_bytes()
        if _is_likely_binary(data):
            raise RuntimeError(
                f"File appears to be binary and cannot be parsed as text: "
                f"{os.path.basename(path)}"
            )
        return data.decode("utf-8", errors="replace")

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _render_page_image(page: Any, scale: int = 3) -> Any:
        """Render a single pypdfium2 page to a PIL Image.

        Uses *scale=3* (~216 DPI) which gives good OCR quality on dense
        layouts without excessive memory use.  Returns ``None`` on
        failure instead of raising.
        """
        try:
            bitmap = page.render(scale=scale)
            return bitmap.to_pil()
        except Exception as e:
            logger.warning("Page render failed: %s", e)
            return None


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _is_likely_binary(data: bytes) -> bool:
    if not data:
        return False
    sample = data[:8192]
    if b"\x00" in sample:
        return True
    non_text = sum(1 for b in sample if (b < 9) or (13 < b < 32))
    return (non_text / len(sample)) > 0.10


def _insert_synthetic_page_markers(text: str, page_size: int = 3000) -> str:
    """Insert ``[PAGE N]`` markers every ~*page_size* chars at paragraph boundaries.

    Used for non-PDF documents (DOCX without page breaks, TXT, MD) so
    they also benefit from page-wise chunking in the vector store.
    """
    if not text or len(text) <= page_size:
        return text

    paragraphs = text.split("\n\n")
    if not paragraphs:
        return text

    out: list[str] = []
    page_num = 1
    current: list[str] = []
    current_len = 0

    for para in paragraphs:
        para_len = len(para)
        if current_len + para_len > page_size and current:
            page_text = "\n\n".join(current).strip()
            if page_text:
                out.append(f"[PAGE {page_num}]\n{page_text}")
                page_num += 1
            current = []
            current_len = 0
        current.append(para)
        current_len += para_len

    if current:
        page_text = "\n\n".join(current).strip()
        if page_text:
            out.append(f"[PAGE {page_num}]\n{page_text}")

    return "\n\n".join(out)


# ---------------------------------------------------------------------------
# Module-level convenience
# ---------------------------------------------------------------------------

_router_instance: ParserRouter | None = None


def get_parser_router(
    default_parser: str = "pdfium",
    complex_parser: str = "pdfium",
    ocr_fallback: str = "rapidocr",
    enable_ocr_fallback: bool = True,
    enable_fast_path: bool = True,
) -> ParserRouter:
    """Return a (lazily-created) singleton ParserRouter."""
    global _router_instance
    if _router_instance is None:
        _router_instance = ParserRouter(
            default_parser=default_parser,
            complex_parser=complex_parser,
            ocr_fallback=ocr_fallback,
            enable_ocr_fallback=enable_ocr_fallback,
            enable_fast_path=enable_fast_path,
        )
    return _router_instance
